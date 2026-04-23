import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def markdown_to_docx(md_file, docx_file):
    document = Document()
    
    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('# '):
            document.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=2)
        
        # Tables (Simple handling)
        elif line.startswith('|'):
            # Skip separator lines like |---|---|
            if '---' in line:
                continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if not cells: continue
            
            # Check if previous element was a table, if not create new
            if len(document.element.body) > 0 and document.element.body[-1].tag.endswith('tbl'):
                table = document.tables[-1]
            else:
                table = document.add_table(rows=0, cols=len(cells))
                table.style = 'Table Grid'
            
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(cells):
                if i < len(row_cells):
                    row_cells[i].text = cell_text
        
        # Images (Placeholder handling)
        elif 'data/spatial_analysis/' in line or 'data/model_results/' in line:
            # Extract path from markdown link or just text
            match = re.search(r'\((.*?)\)', line)
            if match:
                img_path = match.group(1)
            else:
                # Try to find path in raw text
                parts = line.split()
                for p in parts:
                    if 'data/' in p:
                        img_path = p.strip('`').strip('*').strip()
                        break
            
            # Clean path
            img_path = img_path.replace('`', '')
            
            if os.path.exists(img_path):
                try:
                    document.add_picture(img_path, width=Inches(6))
                    last_paragraph = document.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Could not add image {img_path}: {e}")
                    document.add_paragraph(f"[Image: {img_path}]")
            else:
                document.add_paragraph(f"[Image not found: {img_path}]")

        # Bullet points
        elif line.startswith('* ') or line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            document.add_paragraph(text, style='List Number')
            
        # Normal text
        elif line:
            document.add_paragraph(line)

    document.save(docx_file)
    print(f"Word document saved to {docx_file}")

if __name__ == "__main__":
    markdown_to_docx('Analysis_Report.md', 'DeepStreetHeat_Report.docx')
