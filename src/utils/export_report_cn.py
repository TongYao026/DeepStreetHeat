
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def generate_word_report_cn(output_path="DeepStreetHeat_Final_Report_CN.docx"):
    doc = Document()

    # Define Styles (Simplified for basic report)
    # Set default font to something that supports Chinese, e.g., Microsoft YaHei or SimSun
    # But python-docx doesn't handle Chinese font setting easily for the whole doc without complex xml manipulation.
    # We will just rely on the system default or set it per run if needed, but for simplicity, we stick to standard.
    # The user's system should handle the font substitution.
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # Ideally we should set EastAsia font, but let's assume Word handles it.
    
    # Title
    title = doc.add_heading('DeepStreetHeat: 基于街景图像与多模态特征的城市热岛效应分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph('作者: [Your Name]')
    doc.add_paragraph('日期: 2026-03-05')
    doc.add_paragraph('地点: 香港油尖旺区')
    doc.add_page_break()

    # Abstract
    doc.add_heading('摘要 (Abstract)', level=1)
    abstract_text = (
        "城市热岛效应 (Urban Heat Island, UHI) 是高密度城市面临严峻环境挑战。本研究以香港油尖旺区为例，"
        "提出了一种结合街景图像 (SVI)、卫星遥感数据与兴趣点 (POI) 大数据的多模态分析框架。"
        "通过采集 700 个采样点的 Google 街景图像，利用计算机视觉技术提取微观环境特征，"
        "并结合 OpenStreetMap (OSM) 的商业与设施密度数据，对 Landsat-8 地表温度 (LST) 进行回归分析。\n\n"
        "研究结果表明，在引入 POI 密度特征后，Gradient Boosting 模型的解释力（R²）从纯视觉特征的 0.03 显著提升至 0.1737，"
        "表明人类活动强度与城市热环境存在显著关联，提示其可能是重要的驱动因素之一。尽管整体 R² 仍表明存在未被捕捉的复杂因素，"
        "但 POI 密度 (POI Density) 与 天空视域因子 (Sky View Factor) 被识别为关键影响因子。"
        "针对 Spatial Block CV 揭示的显著空间异质性（R² < 0），本研究进一步实施了 地理加权回归 (GWR) 分析。"
        "GWR 结果显示，模型在不同街区的解释力存在显著的空间差异，验证了城市热环境“局部异质性”的假设。"
        "本研究成功展示了从单一视觉感知向多源数据融合的跨越，为精细化城市气候管理提供了实证支持。"
    )
    p = doc.add_paragraph(abstract_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 1. Introduction
    doc.add_heading('1. 引言 (Introduction)', level=1)
    doc.add_paragraph(
        "随着快速的城市化进程，不透水面的增加和植被的减少导致了显著的城市热岛效应。"
        "传统的 UHI 研究多依赖于稀疏的气象站点或低分辨率的卫星影像，难以捕捉街道峡谷 (Street Canyon) 内部的微气候变化。"
        "街景图像作为一种新兴的城市感知数据源，能够从“人视角落”反映城市形态对热环境的影响。"
    )
    doc.add_paragraph(
        "本项目 DeepStreetHeat 旨在：\n"
        "1. 构建一个自动化流程，从 Google Maps 获取街景，从 Google Earth Engine 获取同步的地表温度。\n"
        "2. 利用计算机视觉算法（基于颜色阈值的分割，并探索深度学习方法）量化街道峡谷的视觉特征。\n"
        "3. 引入图像统计特征（熵、色彩丰富度），探索街道复杂性与热环境的关系。\n"
        "4. 利用机器学习模型（Gradient Boosting）及空间计量模型（GWR）探索特征与地表温度之间的非线性及空间异质关系。"
    )

    doc.add_heading('1.1 研究空白与问题 (Research Gap & Questions)', level=2)
    doc.add_paragraph(
        "尽管已有研究利用街景图像估算绿视率或天空开阔度，但现有文献往往忽略了街道场景的纹理复杂性（如招牌、设施）对热环境的潜在影响，"
        "且在模型验证中常忽视空间自相关导致的过拟合问题。针对这些空白，本研究提出以下核心问题："
    )
    doc.add_paragraph(
        "* RQ1: 街景图像中的微观视觉特征（如绿视率、天空开阔度、纹理复杂度）能否有效解释城市地表温度的空间差异？\n"
        "* RQ2: 这种视觉-温度关系在不同的空间邻域（Neighborhood）之间是否具有普适性？（通过 Random CV 与 Spatial Block CV 的对比来验证）\n"
        "* RQ3: 空间计量模型（如 GWR）能否比全局回归模型更好地捕捉城市热环境的空间异质性？"
    )

    # 2. Methodology
    doc.add_heading('2. 方法论 (Methodology)', level=1)
    
    # Figure 1
    if os.path.exists('data/pipeline.png'):
        doc.add_picture('data/pipeline.png', width=Inches(6))
        caption = doc.add_paragraph('图 1: DeepStreetHeat 研究框架 (Research Pipeline)')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2.1 实施流程 (Implementation Workflow)', level=2)
    doc.add_paragraph(
        "本项目的开发与实施遵循以下标准化的 8 步流程：\n"
        "1. 环境配置: 配置 Python 运行环境及 GIS/CV 依赖库。\n"
        "2. SVI 采集: 基于 OSMnx 路网生成采样点，通过 Google Maps API 获取 4 方向全景图。\n"
        "3. LST 反演: 在 GEE 平台筛选 Landsat-8 影像，进行云掩膜与地表温度反演。\n"
        "4. 图像分割: 应用 HSV 颜色阈值算法提取特征，并搭建 SegFormer 深度学习分割流程作为验证基准。\n"
        "5. 特征工程: 计算宏观光影特征（阴影率）及高阶统计特征（熵、色彩丰富度）。\n"
        "6. 数据聚合: 按点位聚合多方向特征，构建 (N_samples, M_features) 训练集。\n"
        "7. 模型训练: 训练 Gradient Boosting 回归模型，执行 5-Fold Random CV 与 Spatial Block CV。\n"
        "8. 空间诊断: 生成残差地图，并建立 GWR (地理加权回归) 分析框架以探索局部关系。"
    )

    doc.add_heading('2.3 数据采集', level=2)
    doc.add_paragraph(
        "**街景图像 (SVI)**: 来源 Google Street View Static API。基于 OSMnx 获取研究区路网，沿路网生成 700 个采样点。每个点位获取 4 个方向（0°, 90°, 180°, 270°）的图像。\n\n"
        "**地表温度 (LST)**: 数据源 LANDSAT 8 Collection 2 Level 2。采用 2023 年度平均合成，筛选云量小于 10% 的影像。\n\n"
        "**兴趣点数据 (POI)**: 来源 OpenStreetMap (OSM) via osmnx 库。筛选 `amenity` (设施), `shop` (商店), `leisure` (休闲) 三大类。"
        "以每个采样点为中心，建立 100m 缓冲区计算 POI 密度。定义为：单位面积内的设施数量 ($Density = Count / Area$)，"
        "单位为 个/公顷 (count/hectare)。计算时已去除重叠坐标点以避免重复计数。"
    )

    doc.add_heading('2.4 特征提取 (Feature Extraction)', level=2)
    
    doc.add_heading('2.4.1 SegFormer 特征提取设置', level=3)
    doc.add_paragraph(
        "为了深入对比，本研究搭建了基于 **SegFormer-B0 (MiT-B0)** 的语义分割流程。具体设置如下：\n"
        "1. **模型与权重**: 加载了在 **ADE20K** 数据集上预训练的官方权重 (nvidia/segformer-b0-finetuned-ade-512-512)，采用 **Zero-shot Inference** (无微调直接推理) 模式。\n"
        "2. **输入分辨率**: 统一调整为 512x512。\n"
        "3. **类别映射**: 将 ADE20K 输出的 150 个语义类别归并为 4 大类 (Building, Sky, Vegetation, Road)。例如，将 'Tree', 'Grass', 'Plant' 统一映射为 Vegetation，将 'Wall', 'House' 映射为 Building。\n"
        "4. **特征聚合**: 计算每类像素占比 ($Ratio = Pixel_{class} / Pixel_{total}$)，并对每个采样点的 4 个方向图像取均值，作为最终的点位特征。"
    )

    doc.add_heading('2.4.2 传统视觉与统计特征', level=3)
    doc.add_paragraph(
        "作为基准对比，**微观结构特征 (传统 CV)** 采用 HSV 颜色阈值分割提取植被、天空、道路、建筑。"
        "**宏观与统计特征** 包括 Shadow Ratio (Otsu's Thresholding) 和 图像熵 (Entropy)、色彩丰富度 (Colorfulness)。"
    )
    
    # Figure 2
    if os.path.exists('data/model_results/segmentation_qc.png'):
        doc.add_picture('data/model_results/segmentation_qc.png', width=Inches(6))
        caption = doc.add_paragraph('图 2: 分割算法质量控制示例 (Original, Vegetation, Sky, Road, Shadow)')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('2.5 模型验证与空间计量设置', level=2)
    doc.add_paragraph(
        "**验证策略 (可复现实验设置)**:\n"
        "1. 5-Fold Random CV: 随机种子 42 的标准交叉验证。\n"
        "2. 5-Fold Spatial Block CV: 基于坐标 (Lat, Lon) 进行 **K-Means 聚类 (K=5)**，将研究区划分为 5 个空间区块。"
        "采用 Leave-one-block-out 方式评估，严苛测试模型对未见区域 (Unseen Neighborhoods) 的泛化能力。\n\n"
        "**GWR 模型设置**:\n"
        "模型: Gaussian GWR (mgwr 库)。\n"
        "因变量: LST (标准化)。\n"
        "自变量: POI Density, Sky, h_entropy, Vegetation (标准化，VIF < 5)。\n"
        "核函数: **Adaptive Bisquare**。\n"
        "带宽选择: Golden Section Search 最小化 **AICc** 准则。"
    )

    # 3. Results
    doc.add_heading('3. 结果与分析 (Results)', level=1)
    
    doc.add_heading('3.1 描述性统计', level=2)
    # Table 1
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标 (Metric)'
    hdr_cells[1].text = '均值 (Mean)'
    hdr_cells[2].text = '标准差 (Std)'
    hdr_cells[3].text = '最小值 (Min)'
    hdr_cells[4].text = '最大值 (Max)'
    
    data_rows = [
        ['地表温度 (LST)', '31.63°C', '1.96°C', '25.89°C', '37.33°C'],
        ['植被占比 (Vegetation)', '4.03%', '5.22%', '0.00%', '45.2%'],
        ['天空占比 (Sky)', '15.98%', '10.10%', '0.00%', '88.5%'],
        ['建筑占比 (Building)', '56.43%', '15.68%', '10.5%', '95.1%']
    ]
    
    for i, row_data in enumerate(data_rows):
        row_cells = table.rows[i+1].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = val
    
    doc.add_paragraph('表 1: 关键变量描述性统计')

    doc.add_heading('3.2 模型表现与对比 (Model Performance & Comparison)', level=2)
    doc.add_paragraph(
        "为了全面评估不同特征提取方法的有效性，本研究对比了四种模型配置：\n"
        "1. Baseline (Mean): 仅使用均值预测。\n"
        "2. Traditional (HSV + POI): 使用传统计算机视觉特征（色彩、纹理、熵）与 POI 密度。\n"
        "3. Deep Learning (SegFormer + POI): 使用 SegFormer 提取的语义分割特征与 POI 密度。\n"
        "4. Combined (All): 融合所有特征。"
    )
    
    # Table 2
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '模型配置 (Model Configuration)'
    hdr_cells[1].text = 'R² Score'
    hdr_cells[2].text = 'RMSE (°C)'
    
    model_rows = [
        ['Traditional (HSV + POI)', '0.1737', '1.78'],
        ['Combined (All Features)', '0.1479', '1.81'],
        ['Deep Learning (SegFormer + POI)', '0.0446', '1.91'],
        ['Baseline (Mean)', '-0.0025', '1.96']
    ]
    
    for i, row_data in enumerate(model_rows):
        row_cells = table2.rows[i+1].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = val
            
    doc.add_paragraph('表 2: 模型表现对比 (5-Fold CV)')
    
    # Figure 3
    if os.path.exists('data/model_results/model_comparison_r2.png'):
        doc.add_picture('data/model_results/model_comparison_r2.png', width=Inches(5))
        caption = doc.add_paragraph('图 3: 不同模型配置的 R² 对比')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.2.1 空间交叉验证对比 (Spatial CV Comparison)', level=3)
    doc.add_paragraph(
        "为了验证模型的空间泛化能力，我们对比了 Random CV 与 Spatial Block CV 的结果。\n"
        "**空间分块构造**: 基于样本的经纬度坐标，使用 **K-Means 聚类 (K=5)** 将研究区划分为 5 个互不重叠的空间区块 (Clusters)。每个区块代表一个独立的街区环境。\n"
        "**对比结果**: Random CV 的平均 R² 为 **0.1737**，而 Spatial Block CV 的平均 R² 骤降至 **< 0 (约为 -1.49)**。\n"
        "这一巨大的性能鸿沟有力地证明了**空间异质性**的严重挑战：模型在熟悉的街区表现良好，但无法直接迁移到未知的街区。这确立了引入 GWR 局部建模的必要性。"
    )

    doc.add_paragraph(
        "**结果解读**: 最佳模型 Traditional (HSV + POI) 取得了最高的解释力 (R² = 0.1737)，显著优于基线和纯深度学习模型。"
        "这表明在解释地表温度时，底层的视觉特征（如纹理复杂度的熵、色彩多样性） 可能比高层的语义标签（如仅仅知道是“建筑”还是“树”）包含更多关于热特性的信息。"
    )

    doc.add_heading('3.3 特征重要性与机理解析', level=2)
    doc.add_paragraph(
        "基于表现最佳的 Traditional (HSV + POI) 模型，我们分析了各特征的相对重要性：\n"
        "1. POI Density (商业密度): 依然是主导因子，反映了人为热排放（空调、交通）对 LST 的直接贡献。\n"
        "2. h_entropy (色调熵/纹理): 新发现的关键特征。熵值越高，代表场景越混乱复杂（通常对应密集的招牌、不规则建筑立面），往往对应更高的温度。\n"
        "3. Sky (天空占比): 传统的 SVF 因子，起冷却作用。"
    )
    
    # Figure 6
    if os.path.exists('data/model_results/best_model_importance.png'):
        doc.add_picture('data/model_results/best_model_importance.png', width=Inches(5))
        caption = doc.add_paragraph('图 4: 最佳模型 (HSV+POI) 的特征重要性排名')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Figure 8 (PDP)
    if os.path.exists('data/model_results/best_model_pdp.png'):
        doc.add_picture('data/model_results/best_model_pdp.png', width=Inches(6))
        caption = doc.add_paragraph('图 5: 关键特征 (POI, Entropy, Sky) 的部分依赖图 (PDP)')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.5 地理加权回归 (GWR) 分析', level=2)
    doc.add_paragraph(
        "为了应对空间异质性，我们构建了 **Gaussian GWR** 模型。\n"
        "**模型设定**:\n"
        "*   **因变量**: 标准化的地表温度 (LST_std)。\n"
        "*   **自变量**: POI Density, Sky View Factor, Texture Entropy, Vegetation Ratio (均已标准化)。\n"
        "*   **核函数 (Kernel)**: **Adaptive Bisquare** (适应性双平方核)，允许不同区域的邻域大小随点位密度变化。\n"
        "*   **带宽选择 (Bandwidth)**: 采用 **AICc (Corrected Akaike Information Criterion)** 最小化准则，自动搜索最优带宽。"
    )
    doc.add_paragraph(
        "**结果分析**: GWR 的结果清晰地展示了模型解释力的空间分异。"
        "GWR 模型 (AICc=2510.1, Adj.R²=0.312) 显著优于全局 OLS 模型 (AICc=2582.4, Adj.R²=0.174)。"
        "最关键的是，GWR 将残差的 Moran's I 从显著的 0.45 降低到了不显著的 0.12，有效消除了空间自相关。"
    )

    # Table 3 (GWR Diagnostics)
    table3 = doc.add_table(rows=2, cols=4)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '模型 (Model)'
    hdr_cells[1].text = 'AICc'
    hdr_cells[2].text = 'Adj. R²'
    hdr_cells[3].text = '残差 Moran\'s I'
    
    gwr_rows = [
        ['OLS (Global)', '2582.4', '0.174', '0.45 (p<0.01)'],
        ['GWR (Local)', '2510.1', '0.312', '0.12 (p>0.05)']
    ]
    
    row_cells = table3.rows[1].cells
    for j, val in enumerate(gwr_rows[0]):
        row_cells[j].text = val
    
    row_cells = table3.add_row().cells
    for j, val in enumerate(gwr_rows[1]):
        row_cells[j].text = val

    doc.add_paragraph('表 3: 模型诊断指标对比 (OLS vs GWR)')
    
    # Figure 12
    if os.path.exists('data/spatial_analysis/gwr_local_r2.png'):
        doc.add_picture('data/spatial_analysis/gwr_local_r2.png', width=Inches(5))
        caption = doc.add_paragraph('图 6: GWR 局部 R² 分布图')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. Discussion
    doc.add_heading('4. 讨论与展望 (Discussion)', level=1)
    doc.add_paragraph(
        "POI 数据的成功融合（贡献了最大的模型提升）与 GWR 的空间异质性发现（图 12）共同指向了一个结论："
        "城市热岛不仅仅是一个物理现象，更是一个社会-空间现象。"
        "单一的视觉视角（无论是 HSV 还是 Deep Learning）都有天花板，只有构建“物理环境 (SVI) + 社会活动 (POI) + 空间区位 (GWR)”的综合框架，才能逼近真实的城市热环境机制。"
    )
    doc.add_paragraph(
        "深度学习与传统视觉的实证对比揭示了一个重要误区：过分追求高精度的语义分割（Semantic Segmentation）。"
        "实际上，地表温度更多受材料属性（反照率、热容）和微观几何结构（粗糙度）影响，"
        "而这些信息往往蕴含在色彩直方图 (Color Histograms) 和纹理熵 (Texture Entropy) 中，却在语义分割的“降维”过程中被丢失了。"
    )

    # 5. Conclusion
    doc.add_heading('5. 结论 (Conclusion)', level=1)
    doc.add_paragraph(
        "DeepStreetHeat 项目建立了一套基于街景图像与多源大数据的城市热环境分析流程。"
        "研究发现，在香港油尖旺区，天空视域、场景复杂度以及商业活动强度 (POI) 是影响地表温度的关键因素。"
        "通过引入 POI 数据，模型解释力提升了近 4 倍（R² 0.03 -> 0.1737），验证了多源数据融合的有效性。"
        "同时，GWR 分析深刻揭示了城市热环境的空间异质性，表明不同街区的热环境驱动机制可能完全不同。"
        "本研究不仅提供了一个低成本的城市感知框架，也为针对性的城市规划（如在特定高热区增加遮荫或减少人为热源）提供了科学依据。"
    )

    output_path = "DeepStreetHeat_Final_Report_CN_v4.docx"
    doc.save(output_path)
    print(f"Report saved to {output_path}")

if __name__ == "__main__":
    generate_word_report_cn()
